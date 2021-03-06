import os
from django.conf import settings
from docx import Document
from io import BytesIO


def create_funding_document(
    funding_body, title, pi_fullname, pi_department, receiver,
    institution_name, template
):
    document = Document(os.path.join(
        settings.BASE_DIR, 
        ('templates/funding/document/' + template))
    )

    document.add_paragraph("Dear {receiver},".format(receiver=receiver))
    document.add_paragraph("")
    document.add_paragraph("")

    paragraph_text = (
        "I can confirm as the PI of the {funding_body} grant "
        "“{title}” "
        "that the research activity is attributable to the Supercomputing "
        "Wales facilities and staff."
    ).format(
        funding_body=funding_body, title=title
    )
    paragraph = document.add_paragraph(paragraph_text)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    document.add_paragraph("")

    paragraph_text = (
        "**Please include a paragraph detailing how Supercomputing "
        "Wales has enabled the research supported by this grant**"
    )
    paragraph = document.add_paragraph(paragraph_text)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    document.add_paragraph("")
    document.add_paragraph("")

    document.add_paragraph("Regards,")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("______________________")
    document.add_paragraph('Prof. {name}'.format(name=pi_fullname))
    if (pi_department):
        document.add_paragraph('{department}'.format(department=pi_department))
    document.add_paragraph(institution_name)
    document.add_paragraph("United Kingdom")
    document.add_paragraph("")

    content = BytesIO()
    document.save(content)
    return content.getvalue()
